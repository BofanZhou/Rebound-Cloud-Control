"""
Generate software-copyright source documents for Rebound Cloud Control.

Outputs:
1) 软著源码汇总.docx
2) 软著源代码提交版_60页.docx
3) 软著源代码汇总_完整版.docx
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable
import math

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


PROJECT_TITLE = "回弹云控管理系统 V2.0"
SUMMARY_DOC_NAME = "软著源码汇总.docx"
SUBMIT_DOC_NAME = "软著源代码提交版_60页.docx"
FULL_DOC_NAME = "软著源代码汇总_完整版.docx"
TARGET_TOTAL_LINES = 3000
LINES_PER_PAGE = 50
TARGET_PAGES = 60
DEFAULT_BALANCED_QUOTA = 1500

KNOWN_PASSWORDS = ("admin123", "maint123", "oper123")


@dataclass(frozen=True)
class SourceFile:
    rel_path: str
    domain: str  # backend | frontend
    lines: list[str]


@dataclass(frozen=True)
class ApiItem:
    method: str
    path: str
    source: str


@dataclass(frozen=True)
class ModelItem:
    name: str
    model_type: str
    source: str


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="生成软著源码汇总与60页源代码提交版文档"
    )
    parser.add_argument(
        "--root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help="项目根目录（默认自动推断）",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help="文档输出目录（默认项目根目录）",
    )
    parser.add_argument(
        "--summary-name",
        default=SUMMARY_DOC_NAME,
        help="源码汇总文档名称",
    )
    parser.add_argument(
        "--submit-name",
        default=SUBMIT_DOC_NAME,
        help="源码提交版文档名称",
    )
    parser.add_argument(
        "--full-name",
        default=FULL_DOC_NAME,
        help="源码完整版文档名称",
    )
    return parser.parse_args()


def iter_core_files(root: Path) -> list[Path]:
    ordered: list[Path] = []
    backend_main = root / "backend" / "main.py"
    if backend_main.exists():
        ordered.append(backend_main)

    backend_app = root / "backend" / "app"
    if backend_app.exists():
        ordered.extend(sorted(backend_app.rglob("*.py")))

    frontend_src = root / "frontend" / "src"
    if frontend_src.exists():
        ordered.extend(sorted(frontend_src.rglob("*.ts")))
        ordered.extend(sorted(frontend_src.rglob("*.vue")))

    deduped: list[Path] = []
    seen: set[str] = set()
    for path in ordered:
        key = str(path.resolve())
        if key in seen:
            continue
        seen.add(key)
        deduped.append(path)
    return deduped


def sanitize_line(line: str, root: Path) -> str:
    cleaned = line

    root_win = str(root.resolve())
    root_posix = root_win.replace("\\", "/")
    cleaned = cleaned.replace(root_win, "<PROJECT_ROOT>")
    cleaned = cleaned.replace(root_posix, "<PROJECT_ROOT>")

    # Known default passwords.
    for pwd in KNOWN_PASSWORDS:
        cleaned = cleaned.replace(pwd, "<PASSWORD>")

    # Common password assignments in code or JSON.
    cleaned = re.sub(
        r'(?i)("password"\s*:\s*")([^"]+)(")',
        r"\1<PASSWORD>\3",
        cleaned,
    )
    cleaned = re.sub(
        r"(?i)(password\s*=\s*['\"])([^'\"]+)(['\"])",
        r"\1<PASSWORD>\3",
        cleaned,
    )

    # Generic Windows absolute paths.
    cleaned = re.sub(r"(?i)\b[A-Z]:\\[^\s\"']+", "<ABS_PATH>", cleaned)
    return cleaned


def read_source_file(path: Path, root: Path) -> SourceFile:
    rel_path = path.relative_to(root).as_posix()
    domain = "backend" if rel_path.startswith("backend/") else "frontend"
    content = path.read_text(encoding="utf-8", errors="ignore")
    lines = content.splitlines()
    sanitized = [sanitize_line(line, root) for line in lines]
    return SourceFile(rel_path=rel_path, domain=domain, lines=sanitized)


def build_domain_lines(files: Iterable[SourceFile], domain: str) -> list[str]:
    flattened: list[str] = []
    for file in files:
        if file.domain != domain:
            continue
        flattened.append(f"# ===== FILE: {file.rel_path} =====")
        if file.lines:
            flattened.extend(file.lines)
        else:
            flattened.append("")
    return flattened


def sample_front_back(lines: list[str], target: int) -> list[str]:
    if target <= 0:
        return []
    if target >= len(lines):
        return list(lines)

    front_count = (target + 1) // 2
    back_count = target - front_count

    prefix = lines[:front_count]
    suffix_source = lines[front_count:]
    suffix = suffix_source[-back_count:] if back_count > 0 else []
    selected = prefix + suffix

    if len(selected) < target:
        middle = lines[front_count : len(lines) - back_count]
        missing = target - len(selected)
        selected.extend(middle[:missing])
    return selected[:target]


def compute_quotas(backend_total: int, frontend_total: int) -> tuple[int, int, str]:
    if (
        backend_total >= DEFAULT_BALANCED_QUOTA
        and frontend_total >= DEFAULT_BALANCED_QUOTA
    ):
        return DEFAULT_BALANCED_QUOTA, DEFAULT_BALANCED_QUOTA, "balanced"

    available_total = backend_total + frontend_total
    if available_total == 0:
        return 0, 0, "empty"

    target = min(TARGET_TOTAL_LINES, available_total)
    backend_quota = round(target * backend_total / available_total)
    frontend_quota = target - backend_quota

    backend_quota = min(backend_quota, backend_total)
    frontend_quota = min(frontend_quota, frontend_total)

    while backend_quota + frontend_quota < target:
        backend_room = backend_total - backend_quota
        frontend_room = frontend_total - frontend_quota
        if backend_room >= frontend_room and backend_room > 0:
            backend_quota += 1
        elif frontend_room > 0:
            frontend_quota += 1
        else:
            break

    return backend_quota, frontend_quota, "proportional"


def build_submission_lines(
    backend_lines: list[str], frontend_lines: list[str]
) -> tuple[list[str], int, int, str]:
    backend_quota, frontend_quota, quota_mode = compute_quotas(
        len(backend_lines), len(frontend_lines)
    )
    selected_backend = sample_front_back(backend_lines, backend_quota)
    selected_frontend = sample_front_back(frontend_lines, frontend_quota)
    merged = selected_backend + selected_frontend

    if len(merged) > TARGET_TOTAL_LINES:
        merged = merged[:TARGET_TOTAL_LINES]
    elif len(merged) < TARGET_TOTAL_LINES:
        merged.extend("# [PADDING]" for _ in range(TARGET_TOTAL_LINES - len(merged)))

    return merged, backend_quota, frontend_quota, quota_mode


def build_full_lines(source_files: list[SourceFile]) -> list[str]:
    lines: list[str] = []
    for source_file in source_files:
        lines.append(f"# ===== FILE: {source_file.rel_path} =====")
        if source_file.lines:
            lines.extend(source_file.lines)
        else:
            lines.append("")
    return lines


def parse_app_version(root: Path) -> str:
    main_path = root / "backend" / "main.py"
    if not main_path.exists():
        return "unknown"
    text = main_path.read_text(encoding="utf-8", errors="ignore")
    match = re.search(r'APP_VERSION\s*=\s*"([^"]+)"', text)
    return match.group(1) if match else "unknown"


def normalize_path(prefix: str, endpoint: str) -> str:
    prefix = (prefix or "").strip()
    endpoint = (endpoint or "").strip()
    if not endpoint:
        endpoint = "/"

    if not prefix:
        path = endpoint
    elif endpoint == "/":
        path = prefix
    else:
        path = f"{prefix.rstrip('/')}/{endpoint.lstrip('/')}"

    if not path.startswith("/"):
        path = "/" + path
    return re.sub(r"/{2,}", "/", path)


def extract_api_items(root: Path) -> list[ApiItem]:
    routers_dir = root / "backend" / "app" / "routers"
    if not routers_dir.exists():
        return []

    items: list[ApiItem] = []
    for path in sorted(routers_dir.glob("*.py")):
        text = path.read_text(encoding="utf-8", errors="ignore")
        rel = path.relative_to(root).as_posix()

        prefix_match = re.search(
            r'APIRouter\([^)]*prefix\s*=\s*["\']([^"\']*)["\']',
            text,
            re.S,
        )
        router_prefix = prefix_match.group(1) if prefix_match else ""

        for match in re.finditer(
            r'@router\.(get|post|put|delete|patch|websocket)\(\s*["\']([^"\']*)["\']',
            text,
        ):
            method = match.group(1).upper()
            endpoint = match.group(2)
            full_path = normalize_path(router_prefix, endpoint)

            if method == "WEBSOCKET":
                items.append(ApiItem(method=method, path=full_path, source=rel))
                continue

            api_path = normalize_path("/api", full_path)
            if full_path != api_path:
                display_path = f"{api_path} (兼容: {full_path})"
            else:
                display_path = api_path
            items.append(ApiItem(method=method, path=display_path, source=rel))
    return items


def extract_model_items(root: Path) -> list[ModelItem]:
    items: list[ModelItem] = []
    backend_app = root / "backend" / "app"
    if backend_app.exists():
        for path in sorted(backend_app.rglob("*.py")):
            text = path.read_text(encoding="utf-8", errors="ignore")
            rel = path.relative_to(root).as_posix()
            for match in re.finditer(r"class\s+(\w+)\(([^)]*)\):", text):
                class_name = match.group(1)
                bases = match.group(2)
                if "BaseModel" in bases:
                    items.append(
                        ModelItem(
                            name=class_name,
                            model_type="Pydantic Model",
                            source=rel,
                        )
                    )
                elif (
                    rel.endswith("backend/app/db/models.py")
                    and re.search(r"\bBase\b", bases)
                    and class_name != "Base"
                ):
                    items.append(
                        ModelItem(
                            name=class_name,
                            model_type="SQLAlchemy Model",
                            source=rel,
                        )
                    )
    return items


def extract_frontend_routes(root: Path) -> list[str]:
    router_path = root / "frontend" / "src" / "router" / "index.ts"
    if not router_path.exists():
        return []
    text = router_path.read_text(encoding="utf-8", errors="ignore")
    routes = re.findall(r"path:\s*'([^']+)'", text)
    unique = []
    seen = set()
    for route in routes:
        if route in seen:
            continue
        seen.add(route)
        unique.append(route)
    return unique


def set_a4_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(18)
        section.right_margin = Mm(18)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)


def apply_mono_style(run) -> None:
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")


def add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{key}：").bold = True
    p.add_run(value)


def add_table_header(table, headers: list[str]) -> None:
    row = table.rows[0]
    for idx, text in enumerate(headers):
        row.cells[idx].text = text


def write_summary_doc(
    output_path: Path,
    source_files: list[SourceFile],
    api_items: list[ApiItem],
    model_items: list[ModelItem],
    frontend_routes: list[str],
    app_version: str,
    backend_quota: int,
    frontend_quota: int,
    quota_mode: str,
) -> None:
    doc = Document()
    set_a4_layout(doc)

    doc.add_heading("软著源码汇总", level=0)
    add_kv(doc, "软件名称", PROJECT_TITLE)
    add_kv(doc, "文档生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    add_kv(doc, "版本号", app_version)
    add_kv(
        doc,
        "源码范围",
        "backend/main.py、backend/app/**.py、frontend/src/**.(ts|vue)",
    )
    add_kv(
        doc,
        "提交版抽取规则",
        f"总计 {TARGET_TOTAL_LINES} 行，{TARGET_PAGES} 页，每页 {LINES_PER_PAGE} 行；"
        f"后端 {backend_quota} 行，前端 {frontend_quota} 行（{quota_mode}）",
    )
    add_kv(
        doc,
        "脱敏策略",
        "仅文档输出脱敏：默认密码替换为 <PASSWORD>，绝对路径替换为 <PROJECT_ROOT>/<ABS_PATH>",
    )

    doc.add_heading("一、技术栈", level=1)
    stack_table = doc.add_table(rows=3, cols=2)
    add_table_header(stack_table, ["层级", "技术"])
    stack_table.rows[1].cells[0].text = "前端"
    stack_table.rows[1].cells[1].text = "Vue 3 / Vite / TypeScript / Pinia / Vue Router / ECharts"
    stack_table.rows[2].cells[0].text = "后端"
    stack_table.rows[2].cells[1].text = "Python / FastAPI / Pydantic / SQLite(SQLAlchemy)"

    doc.add_heading("二、模块-源码映射", level=1)
    mapping = [
        ("前端页面层", "frontend/src/views/*.vue"),
        ("前端路由层", "frontend/src/router/index.ts"),
        ("前端状态管理", "frontend/src/stores/*.ts"),
        ("前端接口封装", "frontend/src/api/index.ts"),
        ("后端API路由层", "backend/app/routers/*.py"),
        ("后端业务服务层", "backend/app/services/*.py"),
        ("后端模型与持久化", "backend/app/models/*.py, backend/app/db/*.py"),
        ("后端应用入口", "backend/main.py"),
    ]
    map_table = doc.add_table(rows=len(mapping) + 1, cols=2)
    add_table_header(map_table, ["功能模块", "对应源码"])
    for idx, (module_name, code_path) in enumerate(mapping, start=1):
        map_table.rows[idx].cells[0].text = module_name
        map_table.rows[idx].cells[1].text = code_path

    doc.add_heading("三、API 清单（自动抽取）", level=1)
    api_table = doc.add_table(rows=max(1, len(api_items)) + 1, cols=3)
    add_table_header(api_table, ["方法", "路径", "来源文件"])
    for idx, api in enumerate(api_items, start=1):
        api_table.rows[idx].cells[0].text = api.method
        api_table.rows[idx].cells[1].text = api.path
        api_table.rows[idx].cells[2].text = api.source

    doc.add_heading("四、数据模型清单（自动抽取）", level=1)
    model_table = doc.add_table(rows=max(1, len(model_items)) + 1, cols=3)
    add_table_header(model_table, ["模型名称", "类型", "来源文件"])
    for idx, model in enumerate(model_items, start=1):
        model_table.rows[idx].cells[0].text = model.name
        model_table.rows[idx].cells[1].text = model.model_type
        model_table.rows[idx].cells[2].text = model.source

    doc.add_heading("五、关键流程", level=1)
    doc.add_paragraph(
        "1. 登录与鉴权流程：前端 LoginView 发起登录请求 -> 后端 auth 路由校验账号/机器 -> "
        "生成 token 与角色信息 -> 路由守卫按角色与机器选择状态放行。"
    )
    doc.add_paragraph(
        "2. 参数推荐与任务执行流程：Dashboard 输入参数 -> /api/recommend 返回补偿参数 -> "
        "/api/device/submit 提交任务 -> 设备模拟器更新状态 -> 历史记录写入数据库。"
    )
    doc.add_paragraph(
        "3. 历史追溯流程：前端 HistoryView 请求 /api/history -> 后端按 machine_id 过滤历史数据 -> "
        "返回统一响应结构并在前端展示。"
    )
    doc.add_paragraph(
        "4. 多机管理流程：用户登录后进入机器选择页 -> /api/machines 获取可用机器 -> "
        "/api/machines/{id}/select 绑定当前机器上下文。"
    )

    doc.add_heading("六、源码行数统计", level=1)
    total_files = len(source_files)
    backend_files = [f for f in source_files if f.domain == "backend"]
    frontend_files = [f for f in source_files if f.domain == "frontend"]
    backend_lines = sum(len(f.lines) for f in backend_files)
    frontend_lines = sum(len(f.lines) for f in frontend_files)
    total_lines = backend_lines + frontend_lines

    add_kv(doc, "核心源码文件数", str(total_files))
    add_kv(doc, "后端源码行数", str(backend_lines))
    add_kv(doc, "前端源码行数", str(frontend_lines))
    add_kv(doc, "核心源码总行数", str(total_lines))

    stat_table = doc.add_table(rows=total_files + 1, cols=2)
    add_table_header(stat_table, ["文件路径", "行数"])
    for idx, source_file in enumerate(source_files, start=1):
        stat_table.rows[idx].cells[0].text = source_file.rel_path
        stat_table.rows[idx].cells[1].text = str(len(source_file.lines))

    doc.add_heading("七、前端路由清单", level=1)
    if frontend_routes:
        for route in frontend_routes:
            doc.add_paragraph(route, style="List Bullet")
    else:
        doc.add_paragraph("未检测到前端路由信息。")

    doc.add_heading("八、运行说明", level=1)
    doc.add_paragraph("1. 生成软著文档：python tools/generate_softcopyright_docs.py")
    doc.add_paragraph("2. 启动后端：cd backend && python main.py")
    doc.add_paragraph("3. 启动前端：cd frontend && npm run dev")
    doc.add_paragraph(
        "4. 说明：该脚本仅生成文档，不修改业务源码，不影响接口与运行行为。"
    )

    doc.save(output_path)


def write_submit_doc(output_path: Path, submit_lines: list[str]) -> None:
    doc = Document()
    set_a4_layout(doc)

    for idx, line in enumerate(submit_lines, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        run = paragraph.add_run(f"{idx:04d} | {line}")
        apply_mono_style(run)

        if idx % LINES_PER_PAGE == 0 and idx < len(submit_lines):
            run.add_break(WD_BREAK.PAGE)

    doc.save(output_path)


def write_full_source_doc(
    output_path: Path,
    full_lines: list[str],
    app_version: str,
) -> None:
    doc = Document()
    set_a4_layout(doc)

    doc.add_heading("软著源代码汇总（完整版）", level=0)
    add_kv(doc, "软件名称", PROJECT_TITLE)
    add_kv(doc, "版本号", app_version)
    add_kv(doc, "文档生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    add_kv(
        doc,
        "源码范围",
        "backend/main.py、backend/app/**.py、frontend/src/**.(ts|vue)",
    )
    add_kv(doc, "总行数", str(len(full_lines)))
    add_kv(
        doc,
        "格式说明",
        f"按源码顺序全量输出，逐行编号；每 {LINES_PER_PAGE} 行分页，便于打印与审阅。",
    )

    doc.add_paragraph(" ")
    for idx, line in enumerate(full_lines, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        run = paragraph.add_run(f"{idx:06d} | {line}")
        apply_mono_style(run)

        if idx % LINES_PER_PAGE == 0 and idx < len(full_lines):
            run.add_break(WD_BREAK.PAGE)

    doc.save(output_path)


def ensure_output_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def main() -> None:
    args = parse_args()
    root = args.root.resolve()
    output_dir = args.output_dir.resolve()
    ensure_output_dir(output_dir)

    core_paths = iter_core_files(root)
    if not core_paths:
        raise SystemExit("未找到核心源码文件，请检查项目目录。")

    source_files = [read_source_file(path, root) for path in core_paths]
    backend_flat = build_domain_lines(source_files, "backend")
    frontend_flat = build_domain_lines(source_files, "frontend")
    full_lines = build_full_lines(source_files)

    submit_lines, backend_quota, frontend_quota, quota_mode = build_submission_lines(
        backend_flat, frontend_flat
    )

    app_version = parse_app_version(root)
    api_items = extract_api_items(root)
    model_items = extract_model_items(root)
    frontend_routes = extract_frontend_routes(root)

    summary_path = output_dir / args.summary_name
    submit_path = output_dir / args.submit_name
    full_path = output_dir / args.full_name

    write_summary_doc(
        output_path=summary_path,
        source_files=source_files,
        api_items=api_items,
        model_items=model_items,
        frontend_routes=frontend_routes,
        app_version=app_version,
        backend_quota=backend_quota,
        frontend_quota=frontend_quota,
        quota_mode=quota_mode,
    )
    write_submit_doc(output_path=submit_path, submit_lines=submit_lines)
    write_full_source_doc(
        output_path=full_path,
        full_lines=full_lines,
        app_version=app_version,
    )

    print("软著文档生成完成：")
    print(f"- {summary_path}")
    print(f"- {submit_path}")
    print(f"- {full_path}")
    print(
        f"提交版规则：{TARGET_TOTAL_LINES} 行 / {TARGET_PAGES} 页 / "
        f"{LINES_PER_PAGE} 行每页，后端 {backend_quota} 行，前端 {frontend_quota} 行。"
    )
    print(
        f"完整版规模：{len(full_lines)} 行 / 约 {math.ceil(len(full_lines) / LINES_PER_PAGE)} 页。"
    )


if __name__ == "__main__":
    main()
